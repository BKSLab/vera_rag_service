from contextlib import contextmanager
from io import BytesIO
from unittest.mock import MagicMock, patch

import docx
import pytest

from app.ingestion.extract import (
    MAX_UPLOAD_SIZE_BYTES,
    TooManyPdfPagesError,
    UnsupportedFileTypeError,
    UploadTooLargeError,
    extract_text_from_upload,
)


def make_docx_bytes(paragraphs: list[str], table_rows: list[tuple[str, str]] | None = None) -> bytes:
    document = docx.Document()
    for paragraph_text in paragraphs:
        document.add_paragraph(paragraph_text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=2)
        for row, (left, right) in zip(table.rows, table_rows, strict=True):
            row.cells[0].text = left
            row.cells[1].text = right
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extract_text_from_upload_decodes_txt():
    result = extract_text_from_upload('doc.txt', 'Текст документа.'.encode())

    assert result == 'Текст документа.'


def test_extract_text_from_upload_decodes_md():
    result = extract_text_from_upload('doc.md', '# Заголовок\nТекст.'.encode())

    assert result == '# Заголовок\nТекст.'


def test_extract_text_from_upload_raises_on_unsupported_extension():
    with pytest.raises(UnsupportedFileTypeError):
        extract_text_from_upload('doc.rtf', b'irrelevant')


def test_extract_text_from_upload_raises_when_file_too_large():
    """ADM-6/ING-7/SEC-4 — лимит размера файла, защита от DoS через
    загрузку чрезмерно большого/специально сконструированного файла."""
    oversized_content = b'a' * (MAX_UPLOAD_SIZE_BYTES + 1)

    with pytest.raises(UploadTooLargeError):
        extract_text_from_upload('doc.txt', oversized_content)


def test_extract_text_from_upload_decodes_docx_paragraphs_preserving_line_breaks():
    """Перенесено из FileTextParser (docx_format_handler.py) — текст параграфов
    в исходном порядке. Важно: переносы строк между параграфами сохраняются
    (не схлопываются в одну строку, как в исходной реализации) — иначе
    `preprocess_document` не найдёт "Статья N" в начале строки."""
    content = make_docx_bytes(['Первый параграф.', 'Статья 21. Квотирование рабочих мест', 'Текст статьи.'])

    result = extract_text_from_upload('doc.docx', content)

    assert result.splitlines() == ['Первый параграф.', 'Статья 21. Квотирование рабочих мест', 'Текст статьи.']


def test_extract_text_from_upload_decodes_docx_table_cells():
    content = make_docx_bytes(['Текст до таблицы.'], table_rows=[('Ячейка 1', 'Ячейка 2')])

    result = extract_text_from_upload('doc.docx', content)

    assert 'Ячейка 1' in result
    assert 'Ячейка 2' in result


def test_extract_text_from_upload_reconstructs_dot_before_superscript_article_suffix():
    """Номер "вставленной" статьи (133¹ = статья 133.1) оформлен в Word как
    обычный ран "Статья 133" + надстрочный ран "1" без разделителя между
    ними — без реконструкции точки текст схлопнулся бы в "Статья 1331",
    неотличимое от другого реального номера статьи (найдено на реальном
    ТК РФ 2026-07-08)."""
    document = docx.Document()
    paragraph = document.add_paragraph('Статья 133')
    superscript_run = paragraph.add_run('1')
    superscript_run.font.superscript = True
    paragraph.add_run('. Установление размера минимальной заработной платы')
    buffer = BytesIO()
    document.save(buffer)

    result = extract_text_from_upload('doc.docx', buffer.getvalue())

    assert result == 'Статья 133.1. Установление размера минимальной заработной платы'


def test_extract_text_from_upload_reconstructs_dot_for_superscript_via_named_style():
    """В реальном ТК РФ (2026-07-08) superscript оформлен не прямым
    свойством рана, а ссылкой на именованный character-стиль (`w:rStyle`) —
    `run.font.superscript` его не видит, нужно разрешать `run.style`."""
    from docx.enum.style import WD_STYLE_TYPE

    document = docx.Document()
    style = document.styles.add_style('SuperscriptStyle', WD_STYLE_TYPE.CHARACTER)
    style.font.superscript = True

    paragraph = document.add_paragraph('Статья 133')
    styled_run = paragraph.add_run('1')
    styled_run.style = style
    paragraph.add_run('. Установление размера минимальной заработной платы')
    buffer = BytesIO()
    document.save(buffer)

    result = extract_text_from_upload('doc.docx', buffer.getvalue())

    assert result == 'Статья 133.1. Установление размера минимальной заработной платы'


def test_extract_text_from_upload_does_not_insert_dot_for_non_adjacent_superscript():
    """Надстрочный текст, не примыкающий к цифре с обеих сторон (например,
    сноска после буквы), не должен получать искусственную точку — вставка
    оправдана только для случая "цифра + надстрочная цифра"."""
    document = docx.Document()
    paragraph = document.add_paragraph('Примечание')
    superscript_run = paragraph.add_run('1')
    superscript_run.font.superscript = True
    buffer = BytesIO()
    document.save(buffer)

    result = extract_text_from_upload('doc.docx', buffer.getvalue())

    assert result == 'Примечание1'


def test_extract_text_from_upload_raises_when_pdf_has_too_many_pages():
    """ADM-6/ING-7/SEC-4 — лимит числа страниц PDF проверяется до
    извлечения текста по каждой странице (самая дорогая операция)."""
    fake_pdf = MagicMock()
    fake_pdf.pages = [MagicMock() for _ in range(2001)]

    @contextmanager
    def fake_open(_buffer):
        yield fake_pdf

    with patch('app.ingestion.extract.pdfplumber.open', fake_open), pytest.raises(TooManyPdfPagesError):
        extract_text_from_upload('doc.pdf', b'%PDF-1.4 fake content')
